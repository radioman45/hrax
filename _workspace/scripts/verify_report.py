# -*- coding: utf-8 -*-
"""생성된 DOCX 내용 검증"""
import sys
import io
import re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

PATH = r"C:\Users\yhlee\Desktop\myprojects\HR AX\survey-feedback\output\leadership_feedback_L001_김민수.docx"

doc = Document(PATH)

print("=== 단락 ===")
for i, p in enumerate(doc.paragraphs):
  if p.text.strip():
    print(f"{i}: {p.text[:150]}")

print("\n=== 표 ===")
for ti, t in enumerate(doc.tables):
  print(f"표 {ti}: {len(t.rows)} 행 x {len(t.columns)} 열")
  for ri, row in enumerate(t.rows):
    cells_text = [c.text.replace(chr(10), ' | ')[:80] for c in row.cells]
    print(f"  R{ri}: {cells_text}")

# 민감정보 재확인
all_text = []
for p in doc.paragraphs:
  all_text.append(p.text)
for t in doc.tables:
  for row in t.rows:
    for c in row.cells:
      for p in c.paragraphs:
        all_text.append(p.text)
full = "\n".join(all_text)

print("\n=== 민감정보 검사 ===")
rid_match = re.findall(r'R\d{3}', full)
# "Report" 는 R+숫자가 아니므로 문제없음
print(f"응답ID(R### 패턴) 매치: {rid_match}")
time_match = re.findall(r'\d{4}-\d{2}-\d{2} \d{2}:\d{2}', full)
print(f"구체 시각(YYYY-MM-DD HH:MM) 매치: {time_match}")
print(f"전체 글자 수: {len(full)}")
