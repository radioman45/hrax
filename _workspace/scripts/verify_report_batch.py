# -*- coding: utf-8 -*-
"""생성된 DOCX 내용 간이 검증 (L002~L006)"""
import sys
import io
import re
import os

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

BASE = r"C:\Users\yhlee\Desktop\myprojects\HR AX\survey-feedback\output"
TARGETS = [
  ("L002", "이지영"),
  ("L003", "박성호"),
  ("L004", "최유진"),
  ("L005", "정현우"),
  ("L006", "한소연"),
]

for leader_id, leader_name in TARGETS:
  path = os.path.join(BASE, f"leadership_feedback_{leader_id}_{leader_name}.docx")
  doc = Document(path)
  all_text = []
  for p in doc.paragraphs:
    all_text.append(p.text)
  for t in doc.tables:
    for row in t.rows:
      for c in row.cells:
        for p in c.paragraphs:
          all_text.append(p.text)
  full = "\n".join(all_text)

  rid = re.findall(r"R\d{3}", full)
  # "Report"는 R+숫자가 아님. 하지만 단어 경계 없이 매치되므로 "R001" 등만 검사됨
  time_match = re.findall(r"\d{4}-\d{2}-\d{2} \d{2}:\d{2}", full)
  para_count = len([p for p in doc.paragraphs if p.text.strip()])
  table_count = len(doc.tables)

  print(f"[{leader_id} {leader_name}]")
  print(f"  파일: {path}")
  print(f"  단락: {para_count}, 표: {table_count}, 전체 글자: {len(full)}")
  print(f"  이름 포함 여부: {leader_name in full}")
  print(f"  팀장명/헤더 확인: {'리더십 Survey Feedback Report' in full}")
  print(f"  응답ID 패턴: {rid}")
  print(f"  구체 시각: {time_match}")
  print()
