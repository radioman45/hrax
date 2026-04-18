# -*- coding: utf-8 -*-
"""
팀장 피드백 리포트 배치 생성 스크립트 (L007~L010)
feedback-report-writer 스킬 규칙에 따라 DOCX 2페이지 생성
색상 팔레트: 헤더 배경 #000000(검정), 섹션 제목 #000000, 배경 #F5F5F5
"""

import json
import os
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 브랜드 상수 (검정/회색 팔레트)
HEADER_COLOR = "000000"   # 헤더 배경 (검정)
ACCENT_COLOR = "000000"   # 섹션 제목, 강조색 (검정)
DARK_GRAY    = "2D2D2D"
LIGHT_BG     = "F5F5F5"   # 섹션 배경 (연한 회색)
FONT_MAIN    = "Noto Sans KR"
FONT_FALLBACK = "나눔고딕"

BASE_DIR = r"C:\Users\yhlee\Desktop\myprojects\HR AX"
INPUT_DIR = os.path.join(BASE_DIR, "_workspace", "01_data")
OUTPUT_DIR = os.path.join(BASE_DIR, "survey-feedback", "output")

TARGETS = [
  ("L007", "오재철"),
  ("L008", "신혜원"),
  ("L009", "윤태경"),
  ("L010", "강서진"),
]


def set_run_font(run, size_pt=10, bold=False, color_hex="2D2D2D", font_name=FONT_MAIN):
  """Run에 폰트 속성 적용 (한글 포함)"""
  run.font.name = font_name
  run.font.size = Pt(size_pt)
  run.font.bold = bold
  r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
  run.font.color.rgb = RGBColor(r, g, b)
  # 한글 폰트 지정 (eastAsia)
  rPr = run._element.get_or_add_rPr()
  rFonts = rPr.find(qn("w:rFonts"))
  if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rPr.append(rFonts)
  rFonts.set(qn("w:eastAsia"), font_name)
  rFonts.set(qn("w:ascii"), font_name)
  rFonts.set(qn("w:hAnsi"), font_name)


def set_paragraph_shading(paragraph, color_hex):
  """단락 배경색 설정"""
  pPr = paragraph._p.get_or_add_pPr()
  shd = OxmlElement("w:shd")
  shd.set(qn("w:val"), "clear")
  shd.set(qn("w:color"), "auto")
  shd.set(qn("w:fill"), color_hex)
  pPr.append(shd)


def set_cell_shading(cell, color_hex):
  """셀 배경색 설정"""
  tcPr = cell._tc.get_or_add_tcPr()
  shd = OxmlElement("w:shd")
  shd.set(qn("w:val"), "clear")
  shd.set(qn("w:color"), "auto")
  shd.set(qn("w:fill"), color_hex)
  tcPr.append(shd)


def set_paragraph_spacing(paragraph, before=0, after=4, line_rule="auto", line=None):
  """단락 간격 설정"""
  pPr = paragraph._p.get_or_add_pPr()
  spacing = pPr.find(qn("w:spacing"))
  if spacing is None:
    spacing = OxmlElement("w:spacing")
    pPr.append(spacing)
  spacing.set(qn("w:before"), str(before * 20))
  spacing.set(qn("w:after"), str(after * 20))
  if line is not None:
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), line_rule)


def score_bar(mean: float, max_score: int = 5) -> str:
  """5점 척도를 텍스트 막대 그래프로 변환"""
  filled = int(round(mean))
  if filled > max_score:
    filled = max_score
  if filled < 0:
    filled = 0
  return "\u2593" * filled + "\u2591" * (max_score - filled)


def add_header(doc, leader_name, team_name, date_range):
  """검정 헤더 바"""
  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.LEFT
  set_paragraph_shading(p, HEADER_COLOR)
  set_paragraph_spacing(p, before=0, after=2)
  run = p.add_run(" 리더십 Survey Feedback Report ")
  set_run_font(run, size_pt=18, bold=True, color_hex="FFFFFF")

  # 부제목
  sub = doc.add_paragraph()
  set_paragraph_spacing(sub, before=2, after=8)
  sub_run = sub.add_run(
    f"{leader_name} 팀장  |  {team_name}  |  응답 기간: {date_range['from']} ~ {date_range['to']}"
  )
  set_run_font(sub_run, size_pt=10, bold=False, color_hex=DARK_GRAY)


def add_section_title(doc, text):
  """섹션 제목 (12pt Bold, 검정)"""
  p = doc.add_paragraph()
  set_paragraph_spacing(p, before=8, after=4)
  run = p.add_run(text)
  set_run_font(run, size_pt=12, bold=True, color_hex=ACCENT_COLOR)


def add_scores_table(doc, scores):
  """섹션 1: Q01~Q10 점수 표 (막대 그래프 포함)"""
  table = doc.add_table(rows=1, cols=4)
  table.style = "Light Grid Accent 1"

  # 헤더 행
  hdr = table.rows[0].cells
  headers = ["문항", "내용", "그래프", "점수"]
  for i, h in enumerate(headers):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0)
    run = p.add_run(h)
    set_run_font(run, size_pt=9, bold=True, color_hex="FFFFFF")
    set_cell_shading(hdr[i], HEADER_COLOR)

  # 열 너비
  col_widths = [Cm(1.5), Cm(8.5), Cm(3.5), Cm(3.5)]
  for i, w in enumerate(col_widths):
    for cell in table.columns[i].cells:
      cell.width = w

  # 데이터 행 (Q01 ~ Q10)
  sorted_qs = sorted(
    [k for k in scores.keys() if k.startswith("Q") and int(k[1:]) <= 10],
    key=lambda x: int(x[1:]),
  )
  for q_id in sorted_qs:
    item = scores[q_id]
    row = table.add_row().cells

    row[0].text = ""
    p0 = row[0].paragraphs[0]
    set_paragraph_spacing(p0, before=0, after=0)
    r0 = p0.add_run(q_id)
    set_run_font(r0, size_pt=9, bold=True, color_hex=DARK_GRAY)

    row[1].text = ""
    p1 = row[1].paragraphs[0]
    set_paragraph_spacing(p1, before=0, after=0)
    r1 = p1.add_run(item["question"])
    set_run_font(r1, size_pt=9, bold=False, color_hex=DARK_GRAY)

    row[2].text = ""
    p2 = row[2].paragraphs[0]
    set_paragraph_spacing(p2, before=0, after=0)
    r2 = p2.add_run(score_bar(item["mean"]))
    set_run_font(r2, size_pt=10, bold=False, color_hex=ACCENT_COLOR)

    row[3].text = ""
    p3 = row[3].paragraphs[0]
    set_paragraph_spacing(p3, before=0, after=0)
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = p3.add_run(f"{item['mean']:.1f} / 5")
    set_run_font(r3, size_pt=9, bold=True, color_hex=DARK_GRAY)


def add_two_column_section(doc, scores, strengths, improvements):
  """섹션 2: 강점/개선 Top 3 2열 레이아웃"""
  table = doc.add_table(rows=2, cols=2)
  table.style = "Table Grid"

  # 1행: 제목
  title_left = table.cell(0, 0)
  title_left.text = ""
  set_cell_shading(title_left, HEADER_COLOR)
  p_tl = title_left.paragraphs[0]
  set_paragraph_spacing(p_tl, before=2, after=2)
  r_tl = p_tl.add_run("  강점 Top 3")
  set_run_font(r_tl, size_pt=11, bold=True, color_hex="FFFFFF")

  title_right = table.cell(0, 1)
  title_right.text = ""
  set_cell_shading(title_right, HEADER_COLOR)
  p_tr = title_right.paragraphs[0]
  set_paragraph_spacing(p_tr, before=2, after=2)
  r_tr = p_tr.add_run("  개선 필요 Top 3")
  set_run_font(r_tr, size_pt=11, bold=True, color_hex="FFFFFF")

  # 2행: 내용
  body_left = table.cell(1, 0)
  body_left.text = ""
  set_cell_shading(body_left, LIGHT_BG)
  first_p = body_left.paragraphs[0]
  set_paragraph_spacing(first_p, before=4, after=2)
  for i, q_id in enumerate(strengths[:3], 1):
    q = scores.get(q_id, {})
    if i == 1:
      p = first_p
    else:
      p = body_left.add_paragraph()
      set_paragraph_spacing(p, before=2, after=2)
    r_num = p.add_run(f"{i}. ")
    set_run_font(r_num, size_pt=10, bold=True, color_hex=ACCENT_COLOR)
    r_q = p.add_run(f"[{q_id}] {q.get('question', '')}")
    set_run_font(r_q, size_pt=10, bold=False, color_hex=DARK_GRAY)
    p2 = body_left.add_paragraph()
    set_paragraph_spacing(p2, before=0, after=4)
    r_s = p2.add_run(f"    평균 {q.get('mean', 0):.1f} / 5")
    set_run_font(r_s, size_pt=9, bold=False, color_hex="666666")

  body_right = table.cell(1, 1)
  body_right.text = ""
  set_cell_shading(body_right, LIGHT_BG)
  first_p_r = body_right.paragraphs[0]
  set_paragraph_spacing(first_p_r, before=4, after=2)
  for i, q_id in enumerate(improvements[:3], 1):
    q = scores.get(q_id, {})
    if i == 1:
      p = first_p_r
    else:
      p = body_right.add_paragraph()
      set_paragraph_spacing(p, before=2, after=2)
    r_num = p.add_run(f"{i}. ")
    set_run_font(r_num, size_pt=10, bold=True, color_hex=ACCENT_COLOR)
    r_q = p.add_run(f"[{q_id}] {q.get('question', '')}")
    set_run_font(r_q, size_pt=10, bold=False, color_hex=DARK_GRAY)
    p2 = body_right.add_paragraph()
    set_paragraph_spacing(p2, before=0, after=4)
    r_s = p2.add_run(f"    평균 {q.get('mean', 0):.1f} / 5")
    set_run_font(r_s, size_pt=9, bold=False, color_hex="666666")


def add_page_break(doc):
  p = doc.add_paragraph()
  run = p.add_run()
  run.add_break(WD_BREAK.PAGE)


def dedupe_preserve_order(items):
  """중복 제거하되 순서 유지"""
  seen = set()
  result = []
  for s in items:
    if s not in seen:
      seen.add(s)
      result.append(s)
  return result


def add_qualitative_section(doc, qualitative, response_count):
  """섹션 3: Q11/Q12 서술형 응답 (익명화, 최대 5개)"""
  if response_count < 3:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=4)
    run = p.add_run("응답 수가 충분하지 않아 개별 의견을 제공할 수 없습니다.")
    set_run_font(run, size_pt=10, bold=False, color_hex=DARK_GRAY)
    return

  labels = [("Q11", "잘하고 있는 점"), ("Q12", "개선이 필요한 점")]
  for q_id, label in labels:
    responses = qualitative.get(q_id, [])
    if not responses:
      continue

    # 소제목
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    r = p.add_run(f"{q_id}. {label}")
    set_run_font(r, size_pt=11, bold=True, color_hex=ACCENT_COLOR)

    # 익명화: 중복 제거 후 최대 5개
    unique_responses = dedupe_preserve_order(responses)[:5]
    for text in unique_responses:
      bp = doc.add_paragraph(style="List Bullet")
      set_paragraph_spacing(bp, before=2, after=2)
      br = bp.add_run(text)
      set_run_font(br, size_pt=10, bold=False, color_hex=DARK_GRAY)


def add_footer_note(doc, date_range):
  """하단 응답 기간 + HR 내부용 표기"""
  sep = doc.add_paragraph()
  set_paragraph_spacing(sep, before=12, after=2)
  sep_run = sep.add_run("\u2500" * 60)
  set_run_font(sep_run, size_pt=8, bold=False, color_hex="CCCCCC")

  p1 = doc.add_paragraph()
  set_paragraph_spacing(p1, before=0, after=2)
  r1 = p1.add_run(f"응답 기간: {date_range['from']} ~ {date_range['to']}")
  set_run_font(r1, size_pt=9, bold=False, color_hex="808080")

  p2 = doc.add_paragraph()
  set_paragraph_spacing(p2, before=0, after=0)
  r2 = p2.add_run("본 리포트는 HR 내부용입니다.")
  set_run_font(r2, size_pt=9, bold=False, color_hex="808080")


def check_sensitive(doc_text: str):
  """민감정보 검사"""
  issues = []
  if re.search(r"R\d{3}", doc_text):
    issues.append("응답ID 패턴 포함 가능성")
  if re.search(r"\d{4}-\d{2}-\d{2} \d{2}:\d{2}", doc_text):
    issues.append("응답 일시 구체값(시분) 포함")
  return issues


def extract_all_text(doc):
  """문서 내 모든 텍스트 수집"""
  texts = []
  for p in doc.paragraphs:
    texts.append(p.text)
  for table in doc.tables:
    for row in table.rows:
      for cell in row.cells:
        for p in cell.paragraphs:
          texts.append(p.text)
  return "\n".join(texts)


def build_report(leader_id: str, leader_name_expected: str):
  input_path = os.path.join(INPUT_DIR, f"{leader_id}_data.json")
  with open(input_path, "r", encoding="utf-8") as f:
    data = json.load(f)

  assert data["leader_id"] == leader_id, f"leader_id mismatch: {data['leader_id']} != {leader_id}"
  assert data["leader_name"] == leader_name_expected, (
    f"leader_name mismatch: {data['leader_name']} != {leader_name_expected}"
  )

  doc = Document()

  # 기본 스타일 폰트
  style = doc.styles["Normal"]
  style.font.name = FONT_MAIN
  style.font.size = Pt(10)
  rPr = style.element.get_or_add_rPr()
  rFonts = rPr.find(qn("w:rFonts"))
  if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rPr.append(rFonts)
  rFonts.set(qn("w:eastAsia"), FONT_MAIN)
  rFonts.set(qn("w:ascii"), FONT_MAIN)
  rFonts.set(qn("w:hAnsi"), FONT_MAIN)

  # A4 여백
  section = doc.sections[0]
  section.page_width = Cm(21)
  section.page_height = Cm(29.7)
  section.top_margin = Cm(1.8)
  section.bottom_margin = Cm(1.8)
  section.left_margin = Cm(2.0)
  section.right_margin = Cm(2.0)

  # 페이지 1
  add_header(doc, data["leader_name"], data["team_name"], data["date_range"])

  add_section_title(doc, "섹션 1. 문항별 점수 요약")
  add_scores_table(doc, data["scores"])

  add_section_title(doc, "섹션 2. 강점 / 개선 필요")
  add_two_column_section(doc, data["scores"], data["strengths"], data["improvements"])

  # 페이지 2
  add_page_break(doc)
  add_section_title(doc, "섹션 3. 구성원 의견 요약")
  add_qualitative_section(doc, data["qualitative"], data["response_count"])
  add_footer_note(doc, data["date_range"])

  # 민감정보 검사
  all_text = extract_all_text(doc)
  issues = check_sensitive(all_text)
  if issues:
    raise RuntimeError(f"[{leader_id}] 민감정보 이슈 발견: {issues}")

  # 저장
  os.makedirs(OUTPUT_DIR, exist_ok=True)
  output_path = os.path.join(
    OUTPUT_DIR, f"leadership_feedback_{leader_id}_{data['leader_name']}.docx"
  )
  doc.save(output_path)
  return output_path, data["response_count"]


def main():
  created = []
  for leader_id, leader_name in TARGETS:
    path, response_count = build_report(leader_id, leader_name)
    print(f"[OK] {leader_id} {leader_name}: {path} (응답 {response_count}건)")
    created.append(path)
  print("\n=== 생성 완료 ===")
  for p in created:
    print(p)


if __name__ == "__main__":
  main()
