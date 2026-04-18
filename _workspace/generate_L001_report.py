# -*- coding: utf-8 -*-
"""
L001 김민수 팀장 리더십 피드백 리포트 생성 (블랙/그레이 톤)
- 헤더 배경: #000000 (검정)
- 섹션 제목 강조: #000000 (검정)
- 섹션 배경: #F5F5F5 (연한 회색)
- SK Orange(#F05023) 미사용
"""
import json
import os
import re
from collections import Counter

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ===== 브랜드 상수 =====
HEADER_COLOR = "000000"      # 헤더 배경: 검정
ACCENT_COLOR = "000000"      # 섹션 제목 강조: 검정
DARK_GRAY = "2D2D2D"
MID_GRAY = "666666"
LIGHT_BG = "F5F5F5"          # 섹션 배경: 연한 회색
BORDER_GRAY = "D0D0D0"
FONT_MAIN = "Noto Sans KR"
FONT_FALLBACK = "맑은 고딕"


# ===== 입출력 경로 =====
BASE_DIR = r"C:\Users\yhlee\Desktop\myprojects\HR AX"
INPUT_PATH = os.path.join(BASE_DIR, "_workspace", "01_data", "L001_data.json")
OUTPUT_PATH = os.path.join(
    BASE_DIR, "survey-feedback", "output", "leadership_feedback_L001_김민수.docx"
)


# ===== 유틸: 폰트 적용 =====
def set_run_font(run, size=10, bold=False, color_hex=DARK_GRAY):
    run.font.name = FONT_MAIN
    # 동아시아 폰트 지정
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_MAIN)
    rFonts.set(qn("w:ascii"), FONT_MAIN)
    rFonts.set(qn("w:hAnsi"), FONT_MAIN)
    run.font.size = Pt(size)
    run.font.bold = bold
    r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)


# ===== 유틸: 단락 배경색 =====
def set_paragraph_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


# ===== 유틸: 셀 배경색 =====
def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


# ===== 유틸: 셀 보더 =====
def set_cell_border(cell, color_hex=BORDER_GRAY, size="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color_hex)
        tcBorders.append(b)
    tcPr.append(tcBorders)


# ===== 유틸: 단락 간격 =====
def set_spacing(paragraph, before=0, after=2, line=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


# ===== 헤더 (검정 배경) =====
def add_header(doc, leader_name, team_name, date_range):
    # 제목 단락: 검정 배경 + 흰색 텍스트
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_shading(p, HEADER_COLOR)
    set_spacing(p, before=2, after=2)

    run = p.add_run("  리더십 Survey Feedback Report")
    set_run_font(run, size=16, bold=True, color_hex="FFFFFF")

    # 부제: 팀장/팀/기간
    sub = doc.add_paragraph()
    set_paragraph_shading(sub, HEADER_COLOR)
    set_spacing(sub, before=0, after=4)
    sub_run = sub.add_run(
        f"  {leader_name} 팀장  |  {team_name}  |  {date_range['from']} ~ {date_range['to']}"
    )
    set_run_font(sub_run, size=10, bold=False, color_hex="E0E0E0")


# ===== 섹션 타이틀 (검정 글씨 + 하단 검정 선) =====
def add_section_title(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=6, after=2)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, color_hex=ACCENT_COLOR)

    # 하단 검정 얇은 선
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), ACCENT_COLOR)
    bottom.set(qn("w:space"), "1")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ===== 섹션 1: 문항별 점수 표 (연한 회색 배경) =====
def score_bar(mean: float, max_score: int = 5) -> str:
    filled = int(round(mean))
    filled = max(0, min(max_score, filled))
    return "■" * filled + "□" * (max_score - filled)


def add_scores_table(doc, scores: dict):
    table = doc.add_table(rows=1 + len(scores), cols=4)
    table.autofit = False

    widths = [Cm(1.2), Cm(9.5), Cm(3.5), Cm(2.8)]
    headers = ["문항", "내용", "응답 분포", "평균"]

    # 헤더 행 (검정 배경 + 흰 글씨)
    hdr = table.rows[0].cells
    for i, (cell, text, w) in enumerate(zip(hdr, headers, widths)):
        cell.width = w
        set_cell_shading(cell, HEADER_COLOR)
        set_cell_border(cell, color_hex=HEADER_COLOR)
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        set_spacing(p, before=0, after=0)
        p.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
        )
        r = p.add_run(text)
        set_run_font(r, size=9, bold=True, color_hex="FFFFFF")

    # 데이터 행
    for row_idx, (q_id, info) in enumerate(scores.items(), start=1):
        row = table.rows[row_idx].cells
        # 지브라: 짝수행 연한 회색
        is_zebra = (row_idx % 2 == 0)
        for i, (cell, w) in enumerate(zip(row, widths)):
            cell.width = w
            if is_zebra:
                set_cell_shading(cell, LIGHT_BG)
            set_cell_border(cell)

        # 문항 ID
        p0 = row[0].paragraphs[0]
        set_spacing(p0, before=0, after=0)
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(q_id)
        set_run_font(r0, size=9, bold=True, color_hex=ACCENT_COLOR)

        # 문항 내용
        p1 = row[1].paragraphs[0]
        set_spacing(p1, before=0, after=0)
        r1 = p1.add_run(info["question"])
        set_run_font(r1, size=9, bold=False, color_hex=DARK_GRAY)

        # 막대
        p2 = row[2].paragraphs[0]
        set_spacing(p2, before=0, after=0)
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r2 = p2.add_run(score_bar(info["mean"]))
        set_run_font(r2, size=10, bold=False, color_hex=ACCENT_COLOR)

        # 평균
        p3 = row[3].paragraphs[0]
        set_spacing(p3, before=0, after=0)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(f"{info['mean']:.2f} / 5")
        set_run_font(r3, size=9, bold=True, color_hex=DARK_GRAY)


# ===== 섹션 2: 강점/개선 2열 =====
def add_strengths_improvements(doc, strengths, improvements, scores):
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    widths = [Cm(8.5), Cm(8.5)]

    # 헤더 행
    headers = ["강점 Top 3", "개선 필요 Top 3"]
    for i, (cell, text) in enumerate(zip(table.rows[0].cells, headers)):
        cell.width = widths[i]
        set_cell_shading(cell, ACCENT_COLOR)
        set_cell_border(cell, color_hex=ACCENT_COLOR)
        p = cell.paragraphs[0]
        p.clear()
        set_spacing(p, before=2, after=2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=10, bold=True, color_hex="FFFFFF")

    # 본문 행
    body_cells = table.rows[1].cells
    for col, q_ids in enumerate([strengths, improvements]):
        cell = body_cells[col]
        cell.width = widths[col]
        set_cell_shading(cell, LIGHT_BG)
        set_cell_border(cell)

        # 첫 단락 초기화
        cell.paragraphs[0].clear()
        first = True
        for idx, q_id in enumerate(q_ids, start=1):
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()
            set_spacing(p, before=2, after=2)
            set_paragraph_shading(p, LIGHT_BG)
            # 번호 + 문항ID
            r_num = p.add_run(f"{idx}. [{q_id}]  ")
            set_run_font(r_num, size=10, bold=True, color_hex=ACCENT_COLOR)
            # 문항 내용
            q_text = scores[q_id]["question"]
            r_q = p.add_run(q_text)
            set_run_font(r_q, size=9, bold=False, color_hex=DARK_GRAY)
            # 평균
            r_m = p.add_run(f"  (평균 {scores[q_id]['mean']:.2f})")
            set_run_font(r_m, size=9, bold=False, color_hex=MID_GRAY)


# ===== 섹션 3: 서술형 =====
def dedupe_preserve_order(items):
    """중복 응답 제거 + 순서 유지. 빈도도 함께 반환."""
    counts = Counter(items)
    seen = set()
    result = []
    for item in items:
        if item in seen:
            continue
        seen.add(item)
        freq = counts[item]
        result.append((item, freq))
    return result


def add_qualitative(doc, qualitative: dict, response_count: int):
    if response_count < 3:
        p = doc.add_paragraph()
        set_spacing(p, before=4, after=4)
        r = p.add_run("응답 수가 충분하지 않아 개별 의견을 제공할 수 없습니다.")
        set_run_font(r, size=10, bold=False, color_hex=DARK_GRAY)
        return

    labels = [("Q11", "잘하고 있는 점"), ("Q12", "개선이 필요한 점")]
    for q_id, label in labels:
        responses = qualitative.get(q_id, [])
        if not responses:
            continue

        # 소제목
        title_p = doc.add_paragraph()
        set_spacing(title_p, before=6, after=2)
        set_paragraph_shading(title_p, LIGHT_BG)
        t_run = title_p.add_run(f"  {q_id}. {label}")
        set_run_font(t_run, size=11, bold=True, color_hex=ACCENT_COLOR)

        # 중복 제거 + 빈도 표시
        deduped = dedupe_preserve_order(responses)
        for text, freq in deduped[:5]:
            bp = doc.add_paragraph()
            set_spacing(bp, before=1, after=1)
            bullet = bp.add_run("  • ")
            set_run_font(bullet, size=10, bold=True, color_hex=ACCENT_COLOR)
            body = bp.add_run(text)
            set_run_font(body, size=10, bold=False, color_hex=DARK_GRAY)
            if freq > 1:
                tag = bp.add_run(f"  (유사 의견 {freq}건)")
                set_run_font(tag, size=9, bold=False, color_hex=MID_GRAY)


# ===== 푸터 (하단 안내) =====
def add_footer_note(doc, date_range):
    # 구분선
    div = doc.add_paragraph()
    set_spacing(div, before=8, after=2)
    pPr = div._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:color"), BORDER_GRAY)
    top.set(qn("w:space"), "1")
    pBdr.append(top)
    pPr.append(pBdr)

    info = doc.add_paragraph()
    set_spacing(info, before=0, after=0)
    r1 = info.add_run(
        f"응답 기간: {date_range['from']} ~ {date_range['to']}    |    본 리포트는 HR 내부용입니다."
    )
    set_run_font(r1, size=9, bold=False, color_hex=MID_GRAY)


# ===== 민감정보 체크 =====
def check_sensitive(doc: Document) -> list:
    issues = []
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text

    if re.search(r"R\d{3}", full_text):
        issues.append("응답ID 패턴(Rxxx) 포함 가능성")
    if re.search(r"\d{4}-\d{2}-\d{2} \d{2}:\d{2}", full_text):
        issues.append("응답 일시 구체값 포함")
    return issues


# ===== 메인 =====
def main():
    with open(INPUT_PATH, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()

    # A4 페이지 설정
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    # 기본 스타일 폰트
    style = doc.styles["Normal"]
    style.font.name = FONT_MAIN
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_MAIN)

    # ===== 페이지 1 =====
    add_header(doc, data["leader_name"], data["team_name"], data["date_range"])

    add_section_title(doc, "섹션 1. 문항별 점수 요약")
    add_scores_table(doc, data["scores"])

    add_section_title(doc, "섹션 2. 강점 / 개선 필요")
    add_strengths_improvements(
        doc, data["strengths"], data["improvements"], data["scores"]
    )

    # 페이지 구분
    doc.add_page_break()

    # ===== 페이지 2 =====
    add_section_title(doc, "섹션 3. 구성원 의견 요약")
    add_qualitative(doc, data["qualitative"], data["response_count"])

    add_footer_note(doc, data["date_range"])

    # 민감정보 체크
    issues = check_sensitive(doc)
    if issues:
        print("[WARN] 민감정보 이슈:", issues)
    else:
        print("[OK] 민감정보 이슈 없음")

    # 저장
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"[SAVED] {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
